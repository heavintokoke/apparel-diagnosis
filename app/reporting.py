from __future__ import annotations

import html
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase import pdfmetrics


def _set_run_font(run, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_paragraph(document: Document, text: str, style: str | None = None):
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(str(text))
    _set_run_font(run, 10)
    return paragraph


def _add_bullets(document: Document, values: list[Any]) -> None:
    for value in values or []:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(str(value))
        _set_run_font(run, 10)


def _evidence_or_gap(item: dict) -> str:
    evidence = item.get("evidence") or []
    if evidence:
        return "\n".join(str(value) for value in evidence[:2])
    return "未在资料中识别到证据，需补充资料或访谈确认。"


def _style_doc(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)
    for name, size, color in [("Heading 1", 17, "1F4E5F"), ("Heading 2", 13, "2C5F2D"), ("Heading 3", 11, "4A5568")]:
        style = document.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def _fill_table(table, rows: list[list[Any]]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, values in enumerate(rows):
        cells = table.rows[row_idx].cells
        for col_idx, value in enumerate(values):
            cell = cells[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 or col_idx in {0, 1, 2} else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(str(value))
            _set_run_font(run, 9, bold=(row_idx == 0))


def export_docx(result: dict, output_path: Path) -> Path:
    document = Document()
    _style_doc(document)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("服装企业全链路经营诊断报告")
    _set_run_font(title_run, 20, True, "1F4E5F")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"生成时间：{result.get('generated_at', '')}｜模式：{result.get('mode', '')}")
    _set_run_font(sub_run, 9, False, "718096")

    document.add_heading("一、公司现状与摘要", level=1)
    _add_paragraph(document, result.get("company_snapshot", ""))
    _add_paragraph(document, result.get("executive_summary", ""))
    if result.get("ai_status"):
        _add_paragraph(document, result["ai_status"])

    document.add_heading("二、核心问题", level=1)
    _add_bullets(document, result.get("core_findings", []))

    document.add_heading("三、根因分析", level=1)
    _add_bullets(document, result.get("root_causes", []))

    document.add_heading("四、主线判断", level=1)
    rows = [["主线", "得分", "风险", "判断", "下一步动作"]]
    for line in result.get("main_lines", []):
        rows.append(
            [
                line.get("name", ""),
                line.get("score", ""),
                line.get("risk_label", ""),
                line.get("judgment", ""),
                "\n".join(line.get("next_actions", [])[:3]),
            ]
        )
    if len(rows) > 1:
        table = document.add_table(rows=len(rows), cols=5)
        _fill_table(table, rows)

    document.add_heading("五、全链路流程图", level=1)
    rows = [["流程", "链路步骤", "风险", "改善重点"]]
    for flow in result.get("flow_charts", []):
        rows.append([flow.get("name", ""), flow.get("content", ""), flow.get("risk_label", ""), flow.get("improvement_focus", "")])
    if len(rows) > 1:
        table = document.add_table(rows=len(rows), cols=4)
        _fill_table(table, rows)

    document.add_heading("六、部门诊断问题清单", level=1)
    rows = [["部门", "建议责任人", "风险", "关联板块", "下一步动作"]]
    for department in result.get("department_issues", []):
        rows.append(
            [
                department.get("department", ""),
                department.get("owner", ""),
                department.get("risk_label", ""),
                "\n".join(department.get("linked_sections", [])[:5]),
                department.get("next_action", ""),
            ]
        )
    if len(rows) > 1:
        table = document.add_table(rows=len(rows), cols=5)
        _fill_table(table, rows)

    document.add_heading("七、16 个板块诊断结果", level=1)
    section_rows = [["板块", "得分", "风险", "诊断结论"]]
    for section in result.get("sections", []):
        section_rows.append([section["title"], section["score"], section["risk_label"], section["finding"]])
    table = document.add_table(rows=len(section_rows), cols=4)
    _fill_table(table, section_rows)

    document.add_heading("八、诊断项证据明细", level=1)
    detail_rows = [["板块", "诊断项", "状态", "证据 / 缺口", "建议动作"]]
    for section in result.get("sections", []):
        for item in section.get("items", []):
            detail_rows.append(
                [
                    section.get("title", ""),
                    item.get("name", ""),
                    item.get("status_label", ""),
                    _evidence_or_gap(item),
                    item.get("improvement_action", ""),
                ]
            )
    if len(detail_rows) > 1:
        table = document.add_table(rows=len(detail_rows), cols=5)
        _fill_table(table, detail_rows)

    document.add_heading("九、关键数据缺失清单", level=1)
    missing = result.get("missing_data", [])
    if missing:
        rows = [["板块", "缺失指标", "意义", "建议责任人"]]
        rows.extend([[item["section"], item["metric"], item["meaning"], item["suggested_owner"]] for item in missing])
        table = document.add_table(rows=len(rows), cols=4)
        _fill_table(table, rows)
    else:
        _add_paragraph(document, "当前资料未识别到关键指标缺失。")

    document.add_heading("十、改善路径", level=1)
    document.add_heading("短期改善", level=2)
    _add_bullets(document, result.get("short_term_improvements", []))
    document.add_heading("中期系统建设", level=2)
    _add_bullets(document, result.get("medium_term_system_build", []))
    document.add_heading("长期增长路径", level=2)
    _add_bullets(document, result.get("long_term_growth_path", []))

    document.add_heading("十一、90 天经营改善计划", level=1)
    plan = result.get("ninety_day_plan", [])
    rows = [["优先级", "问题", "目标", "动作", "负责人", "时间", "检查指标", "复盘机制"]]
    for idx, item in enumerate(plan, start=1):
        rows.append(
            [
                item.get("priority", idx),
                item.get("problem", ""),
                item.get("goal", ""),
                item.get("action", ""),
                item.get("owner", ""),
                item.get("timeframe", ""),
                item.get("check_metric", ""),
                item.get("review_mechanism", ""),
            ]
        )
    table = document.add_table(rows=len(rows), cols=8)
    _fill_table(table, rows)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def _p(text: Any, style: ParagraphStyle):
    return Paragraph(html.escape(str(text or "")).replace("\n", "<br/>"), style)


def export_pdf(result: dict, output_path: Path) -> Path:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ChineseTitle",
        parent=styles["Title"],
        fontName="STSong-Light",
        fontSize=20,
        leading=26,
        textColor=colors.HexColor("#1F4E5F"),
        alignment=1,
    )
    h1 = ParagraphStyle("ChineseH1", parent=styles["Heading1"], fontName="STSong-Light", fontSize=14, leading=20, textColor=colors.HexColor("#1F4E5F"))
    body = ParagraphStyle("ChineseBody", parent=styles["BodyText"], fontName="STSong-Light", fontSize=9.5, leading=15)
    small = ParagraphStyle("ChineseSmall", parent=body, fontSize=8, leading=12)
    tiny = ParagraphStyle("ChineseTiny", parent=body, fontSize=7, leading=10)

    story: list[Any] = [_p("服装企业全链路经营诊断报告", title_style), Spacer(1, 6 * mm)]
    story.append(_p(f"生成时间：{result.get('generated_at', '')}｜模式：{result.get('mode', '')}", small))
    story.append(Spacer(1, 5 * mm))

    for heading, values in [
        ("一、公司现状与摘要", [result.get("company_snapshot", ""), result.get("executive_summary", ""), result.get("ai_status", "")]),
        ("二、核心问题", result.get("core_findings", [])),
        ("三、根因分析", result.get("root_causes", [])),
    ]:
        story.append(_p(heading, h1))
        for value in values:
            if value:
                story.append(_p(f"• {value}" if isinstance(values, list) and heading != "一、公司现状与摘要" else value, body))
        story.append(Spacer(1, 3 * mm))

    story.append(_p("四、主线判断", h1))
    rows = [[_p("主线", small), _p("得分", small), _p("风险", small), _p("判断", small), _p("下一步动作", small)]]
    for line in result.get("main_lines", []):
        rows.append([
            _p(line.get("name", ""), small),
            line.get("score", ""),
            _p(line.get("risk_label", ""), small),
            _p(line.get("judgment", ""), small),
            _p("；".join(line.get("next_actions", [])[:3]), small),
        ])
    if len(rows) > 1:
        table = Table(rows, colWidths=[24 * mm, 13 * mm, 18 * mm, 58 * mm, 57 * mm], repeatRows=1)
        table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E0")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E6F0EE")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
        story.extend([table, Spacer(1, 5 * mm)])

    story.append(_p("五、全链路流程图", h1))
    rows = [[_p("流程", small), _p("链路步骤", small), _p("风险", small), _p("改善重点", small)]]
    for flow in result.get("flow_charts", []):
        rows.append([_p(flow.get("name", ""), small), _p(flow.get("content", ""), small), _p(flow.get("risk_label", ""), small), _p(flow.get("improvement_focus", ""), small)])
    if len(rows) > 1:
        table = Table(rows, colWidths=[32 * mm, 72 * mm, 20 * mm, 46 * mm], repeatRows=1)
        table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E0")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EEF4F8")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
        story.extend([table, Spacer(1, 5 * mm)])

    story.append(_p("六、部门诊断问题清单", h1))
    rows = [[_p("部门", small), _p("负责人", small), _p("风险", small), _p("关联板块", small), _p("下一步动作", small)]]
    for department in result.get("department_issues", []):
        rows.append([
            _p(department.get("department", ""), small),
            _p(department.get("owner", ""), small),
            _p(department.get("risk_label", ""), small),
            _p("；".join(department.get("linked_sections", [])[:5]), tiny),
            _p(department.get("next_action", ""), tiny),
        ])
    if len(rows) > 1:
        table = Table(rows, colWidths=[18 * mm, 24 * mm, 20 * mm, 50 * mm, 58 * mm], repeatRows=1)
        table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E0")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#FFF4D6")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
        story.extend([table, Spacer(1, 5 * mm)])

    story.append(_p("七、16 个板块诊断结果", h1))
    rows = [[_p("板块", small), _p("得分", small), _p("风险", small), _p("诊断结论", small)]]
    for section in result.get("sections", []):
        rows.append([_p(section["title"], small), section["score"], _p(section["risk_label"], small), _p(section["finding"], small)])
    table = Table(rows, colWidths=[38 * mm, 16 * mm, 20 * mm, 96 * mm], repeatRows=1)
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E0")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E6F0EE")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
    story.extend([table, Spacer(1, 5 * mm)])

    story.append(_p("八、诊断项证据明细", h1))
    rows = [[_p("板块", tiny), _p("诊断项", tiny), _p("状态", tiny), _p("证据 / 缺口", tiny), _p("建议动作", tiny)]]
    for section in result.get("sections", []):
        for item in section.get("items", []):
            rows.append([
                _p(section.get("title", ""), tiny),
                _p(item.get("name", ""), tiny),
                _p(item.get("status_label", ""), tiny),
                _p(_evidence_or_gap(item), tiny),
                _p(item.get("improvement_action", ""), tiny),
            ])
    if len(rows) > 1:
        table = Table(rows, colWidths=[30 * mm, 25 * mm, 18 * mm, 50 * mm, 47 * mm], repeatRows=1)
        table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.2, colors.HexColor("#CBD5E0")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F7F5")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
        story.extend([table, Spacer(1, 5 * mm)])

    story.append(_p("九、关键数据缺失清单", h1))
    missing = result.get("missing_data", [])
    if missing:
        rows = [[_p("板块", small), _p("缺失指标", small), _p("意义", small), _p("建议责任人", small)]]
        for item in missing:
            rows.append([_p(item["section"], small), _p(item["metric"], small), _p(item["meaning"], small), _p(item["suggested_owner"], small)])
        table = Table(rows, colWidths=[38 * mm, 28 * mm, 76 * mm, 28 * mm], repeatRows=1)
        table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E0")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#FFF4D6")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
        story.extend([table, Spacer(1, 5 * mm)])

    story.append(_p("十、90 天经营改善计划", h1))
    rows = [[_p("优先级", small), _p("问题", small), _p("动作", small), _p("负责人", small), _p("检查指标", small), _p("复盘机制", small)]]
    for idx, item in enumerate(result.get("ninety_day_plan", []), start=1):
        rows.append([
            item.get("priority", idx),
            _p(item.get("problem", ""), small),
            _p(item.get("action", ""), small),
            _p(item.get("owner", ""), small),
            _p(item.get("check_metric", ""), small),
            _p(item.get("review_mechanism", ""), small),
        ])
    table = Table(rows, colWidths=[13 * mm, 38 * mm, 46 * mm, 22 * mm, 24 * mm, 32 * mm], repeatRows=1)
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E0")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E6F0EE")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
    story.append(table)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(str(output_path), pagesize=A4, rightMargin=15 * mm, leftMargin=15 * mm, topMargin=15 * mm, bottomMargin=15 * mm)
    doc.build(story)
    return output_path
